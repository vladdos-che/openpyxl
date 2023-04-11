import docx
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

doc = docx.Document('text.docx')
text = doc.paragraphs[0].runs[1].text
print(text)


new_doc = docx.Document()

style = new_doc.styles.add_style('Down', WD_STYLE_TYPE.PARAGRAPH)
style.font.name = 'Bahnschrift'
style.font.size = Pt(20)

str_01 = new_doc.add_paragraph('Здравствуй, мир! ')
str_01.add_run(f'Hi, {text}!')

for paragraph in new_doc.paragraphs:
    paragraph.style = 'Down'

new_doc.save("new_text.docx")
